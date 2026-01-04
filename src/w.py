from __future__ import annotations
from typing import Dict, Tuple
from docx import Document

from models import ApplicationRow


def _norm(s: str) -> str:
    return (s or "").strip()


def parse_application_docx(path: str) -> ApplicationRow:
    doc = Document(path)
    row = ApplicationRow()

    # 1) Parse tables (common in forms)
    # Strategy: treat table rows as "Label | Value"
    for table in doc.tables:
        for r in table.rows:
            cells = [c.text.strip() for c in r.cells]
            if len(cells) < 2:
                continue
            label = _norm(cells[0]).lower()
            value = _norm(cells[1])

            if "name" == label or "student name" in label:
                row.student_name = value
            elif "pronoun" in label:
                row.pronouns = value
            elif label == "email" or "email address" in label:
                row.student_email = value
            elif "faculty" in label:
                row.faculty = value
            elif "major" in label:
                row.major_area = value
            elif "year" in label:
                row.year_of_study = value

    # 2) Checkboxes (often appear in paragraphs or table cells)
    full_text = "\n".join([p.text for p in doc.paragraphs]).lower()

    # crude but works if the form uses these words near checkbox markers
    row.first_generation = "first generation" in full_text and ("☑" in full_text or "checked" in full_text)
    row.indigenous = "indigenous" in full_text and "☑" in full_text
    row.racialized = "racialized" in full_text and "☑" in full_text
    row.lgbtq2si = "lgbtq" in full_text and "☑" in full_text
    row.disability = "disability" in full_text and "☑" in full_text
    row.international = "international" in full_text and "☑" in full_text

    # 3) Essay capture by headings
    row.research_statement = _extract_section(doc, ["research statement", "research essay", "research"])
    row.leadership_statement = _extract_section(doc, ["leadership statement", "leadership essay", "leadership"])

    return row


def _extract_section(doc: Document, heading_keywords: list[str]) -> str:
    paras = [p.text.strip() for p in doc.paragraphs]
    # find heading
    start = None
    for i, t in enumerate(paras):
        tl = t.lower()
        if any(k in tl for k in heading_keywords) and len(t) <= 80:
            start = i + 1
            break
    if start is None:
        return ""

    # collect until next "heading-like" short line
    out = []
    for j in range(start, len(paras)):
        t = paras[j]
        if not t:
            continue
        if len(t) <= 80 and t.endswith(":"):
            break
        # stop if another section header appears
        tl = t.lower()
        if any(k in tl for k in ["metadata", "signature", "leadership", "research statement", "academic info", "student info"]):
            # heuristic: if looks like a section switch
            if len(t) <= 80:
                break
        out.append(t)

    return "\n".join(out).strip()
